"""
驗證 DOCX 產出：
L1 結構（python-docx 唯讀）
L2 視覺（Word COM → PDF → PNG）
"""
import sys
from pathlib import Path
from docx import Document


def verify_structure(docx_path: Path):
    """L1：結構驗證"""
    print(f'--- L1 結構驗證：{docx_path.name} ---')
    doc = Document(str(docx_path))

    # Metadata
    cp = doc.core_properties
    print(f'  Title:   {cp.title}')
    print(f'  Author:  {cp.author}')
    print(f'  Subject: {cp.subject}')

    # 段落與表格
    para_count = len(doc.paragraphs)
    table_count = len(doc.tables)
    print(f'  段落數：{para_count}')
    print(f'  表格數：{table_count}')

    # CJK 字型抽樣檢查（前 20 個 run）
    cjk_font_issues = 0
    checked = 0
    for para in doc.paragraphs[:50]:
        for run in para.runs:
            if not run.text.strip():
                continue
            checked += 1
            # 檢查是否有 eastAsia 字型設定
            rPr = run._element.find(
                '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}rPr'
            )
            if rPr is None:
                cjk_font_issues += 1
                continue
            rFonts = rPr.find(
                '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}rFonts'
            )
            if rFonts is None:
                cjk_font_issues += 1
                continue
            east_asia = rFonts.get(
                '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}eastAsia'
            )
            if not east_asia:
                cjk_font_issues += 1

            if checked >= 20:
                break
        if checked >= 20:
            break

    print(f'  CJK 字型檢查：{checked} 個 run 中 {cjk_font_issues} 個缺少 eastAsia 設定')
    if cjk_font_issues == 0:
        print('  ✅ CJK 雙字型設定正確')
    else:
        print('  ⚠️ 部分 run 缺少 eastAsia——可能是中性文字（無中文）')

    # 表格大小檢查
    for i, table in enumerate(doc.tables):
        rows = len(table.rows)
        cols = len(table.columns) if table.rows else 0
        print(f'  表格 #{i+1}: {rows} rows × {cols} cols')

    return True


def export_pdf(docx_path: Path, pdf_path: Path):
    """L2：Word COM 轉 PDF"""
    print(f'--- L2 視覺驗證：匯出 PDF ---')
    import win32com.client
    import pythoncom

    pythoncom.CoInitialize()
    word = win32com.client.Dispatch('Word.Application')
    word.Visible = False
    try:
        doc = word.Documents.Open(str(docx_path))
        # wdExportFormatPDF = 17
        doc.ExportAsFixedFormat(
            OutputFileName=str(pdf_path),
            ExportFormat=17,
            OpenAfterExport=False,
            OptimizeFor=0,           # Print quality
            Range=0,                 # Entire document
            IncludeDocProps=True,
            CreateBookmarks=1,       # Use headings
        )
        doc.Close(SaveChanges=False)
        print(f'  ✅ PDF 產出：{pdf_path}')
        print(f'     檔案大小：{pdf_path.stat().st_size / 1024:.1f} KB')
    finally:
        word.Quit()
